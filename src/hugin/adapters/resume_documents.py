from __future__ import annotations

import hashlib
from pathlib import Path
from zipfile import BadZipFile

from docx import Document as WordDocument
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from hugin.domain.resumes import ResumeDocument, ResumeSourceType


class ResumeDocumentError(ValueError):
    pass


class ResumeDocumentReader:
    def __init__(
        self,
        *,
        max_size_bytes: int = 20 * 1024 * 1024,
        min_text_chars: int = 200,
    ) -> None:
        self._max_size_bytes = max_size_bytes
        self._min_text_chars = min_text_chars

    def read(self, source: Path) -> ResumeDocument:
        path = source.expanduser().resolve(strict=True)
        if not path.is_file():
            raise ResumeDocumentError("Путь резюме не является файлом")

        size_bytes = path.stat().st_size
        if size_bytes == 0:
            raise ResumeDocumentError("Файл резюме пуст")
        if size_bytes > self._max_size_bytes:
            raise ResumeDocumentError("Файл резюме превышает допустимый размер")

        suffix = path.suffix.casefold()
        page_count: int | None
        if suffix == ".pdf":
            source_type = ResumeSourceType.PDF
            text, page_count = self._read_pdf(path)
        elif suffix == ".docx":
            source_type = ResumeSourceType.DOCX
            text, page_count = self._read_docx(path)
        else:
            raise ResumeDocumentError("Поддерживаются только файлы PDF и DOCX")

        normalized = self._normalize_text(text)
        if len(normalized) < self._min_text_chars:
            raise ResumeDocumentError(
                "В документе нет пригодного текстового слоя; требуется распознавание"
            )
        if "\ufffd" in normalized or "\x00" in normalized:
            raise ResumeDocumentError("Текст документа повреждён")

        return ResumeDocument(
            source_path=path,
            original_name=path.name,
            source_type=source_type,
            sha256=self._sha256(path),
            size_bytes=size_bytes,
            page_count=page_count,
            text=normalized,
        )

    @staticmethod
    def _read_pdf(path: Path) -> tuple[str, int]:
        try:
            reader = PdfReader(path)
            if reader.is_encrypted:
                raise ResumeDocumentError("Защищённый PDF нельзя импортировать")
            pages = [(page.extract_text() or "") for page in reader.pages]
        except ResumeDocumentError:
            raise
        except (OSError, PdfReadError, ValueError) as error:
            raise ResumeDocumentError("Не удалось прочитать PDF") from error
        return "\n".join(pages), len(pages)

    @staticmethod
    def _read_docx(path: Path) -> tuple[str, None]:
        try:
            document = WordDocument(str(path))
        except (BadZipFile, OSError, PackageNotFoundError, ValueError) as error:
            raise ResumeDocumentError("Не удалось прочитать DOCX") from error

        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts), None

    @staticmethod
    def _normalize_text(text: str) -> str:
        lines = (" ".join(line.split()) for line in text.splitlines())
        return "\n".join(line for line in lines if line)

    @staticmethod
    def _sha256(path: Path) -> str:
        digest = hashlib.sha256()
        with path.open("rb") as source:
            for chunk in iter(lambda: source.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()
