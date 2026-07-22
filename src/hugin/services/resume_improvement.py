from __future__ import annotations

import hashlib
import json
import os
import re
from collections import Counter
from collections.abc import Callable
from dataclasses import asdict, dataclass
from datetime import UTC, datetime
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Protocol

from docx import Document
from docx.shared import Pt
from sqlalchemy import select
from sqlalchemy.orm import Session

from hugin.database.models import (
    CandidateProfileModel,
    CareerDirectionModel,
    DirectionVacancyModel,
    ResumeModel,
    VacancyModel,
)
from hugin.services.resume_prompts import (
    QUESTION_PROMPT_VERSION,
    REWRITE_PROMPT_VERSION,
    SYSTEM_PROMPT,
    ResumeBlockKind,
    ResumePromptContext,
    ResumeQuestionAnswer,
    ResumeQuestionAssessment,
    build_questions_prompt,
    build_rewrite_prompt,
    parse_question_assessments,
    select_missing_questions,
)


class ResumeTextModel(Protocol):
    @property
    def model_name(self) -> str: ...

    def complete(self, system_prompt: str, user_prompt: str) -> str: ...


@dataclass(frozen=True, slots=True)
class ResumeNarrativeBlock:
    index: int
    kind: ResumeBlockKind
    label: str
    start_line: int
    end_line: int
    source_text: str


@dataclass(frozen=True, slots=True)
class ResumeStructure:
    lines: tuple[str, ...]
    blocks: tuple[ResumeNarrativeBlock, ...]


@dataclass(frozen=True, slots=True)
class ImprovedResumeBlock:
    index: int
    kind: ResumeBlockKind
    label: str
    source_text: str
    improved_text: str
    questions: tuple[str, ...]
    answers: tuple[str, ...]


@dataclass(frozen=True, slots=True)
class ResumeImprovementResult:
    resume_id: int
    target_role: str
    model_name: str
    draft_path: Path
    report_path: Path
    blocks: tuple[ImprovedResumeBlock, ...]
    source_unchanged: bool


AnswerProvider = Callable[[ResumeNarrativeBlock, str], str]


class ResumeBlockExtractor:
    _months = (
        "Январь",
        "Февраль",
        "Март",
        "Апрель",
        "Май",
        "Июнь",
        "Июль",
        "Август",
        "Сентябрь",
        "Октябрь",
        "Ноябрь",
        "Декабрь",
    )
    _date_start = re.compile(rf"^(?:{'|'.join(_months)})\s+\d{{4}}\s*[—-]", re.IGNORECASE)
    _date_line = re.compile(
        rf"^(?:{'|'.join(_months)})\s+\d{{4}}(?:\s*[—-])?$",
        re.IGNORECASE,
    )
    _duration = re.compile(r"^(?:\d+\s+)?(?:год|года|лет|месяц|месяца|месяцев)\b", re.IGNORECASE)
    _project = re.compile(r"^Проект\s+(.+?):$", re.IGNORECASE)
    _project_item = re.compile(r"^-\s+(.{2,100}?)\s+[—-]\s+.+")
    _top_sections = (
        "Образование",
        "Повышение квалификации",
        "Навыки",
        "Знание языков",
        "Опыт вождения",
        "Дополнительная информация",
    )

    def extract(self, source_text: str) -> ResumeStructure:
        lines = self._clean_lines(source_text)
        work_start = self._find_work_start(lines)
        work_end = self._find_work_end(lines, work_start)
        entry_starts = [
            index
            for index in range(work_start, work_end)
            if self._date_start.search(lines[index]) is not None
        ]
        if not entry_starts:
            raise ValueError("В разделе опыта не найдены отдельные места работы")

        blocks: list[ResumeNarrativeBlock] = []
        for entry_number, entry_start in enumerate(entry_starts, start=1):
            entry_end = entry_starts[entry_number] if entry_number < len(entry_starts) else work_end
            narrative_start = self._find_narrative_start(lines, entry_start, entry_end)
            project_spans = self._project_spans(lines, narrative_start, entry_end)
            first_project_marker = min(
                (marker for marker, _, _, _ in project_spans),
                default=entry_end,
            )
            if narrative_start < first_project_marker:
                blocks.append(
                    self._block(
                        blocks,
                        ResumeBlockKind.WORK_EXPERIENCE,
                        self._work_label(lines, entry_start, narrative_start, entry_number),
                        narrative_start,
                        first_project_marker,
                        lines,
                    )
                )
            for _, project_start, project_end, label in project_spans:
                blocks.append(
                    self._block(
                        blocks,
                        ResumeBlockKind.PROJECT,
                        label,
                        project_start,
                        project_end,
                        lines,
                    )
                )

        if not blocks:
            raise ValueError("В опыте работы не найдены содержательные блоки")
        return ResumeStructure(tuple(lines), tuple(blocks))

    def assemble(
        self,
        structure: ResumeStructure,
        improved_blocks: tuple[ImprovedResumeBlock, ...],
    ) -> str:
        replacements = {block.index: block.improved_text for block in improved_blocks}
        if set(replacements) != {block.index for block in structure.blocks}:
            raise ValueError("Для сборки должны быть готовы все блоки резюме")

        lines = list(structure.lines)
        for block in sorted(structure.blocks, key=lambda item: item.start_line, reverse=True):
            replacement = replacements[block.index].splitlines()
            lines[block.start_line : block.end_line] = replacement
        return "\n".join(lines).strip() + "\n"

    @staticmethod
    def _clean_lines(source_text: str) -> list[str]:
        lines: list[str] = []
        for raw_line in source_text.splitlines():
            line = " ".join(raw_line.replace("\u00a0", " ").split())
            if not line or "Резюме обновлено" in line:
                continue
            lines.append(line)
        if not lines:
            raise ValueError("Текст резюме пуст")
        return lines

    @staticmethod
    def _find_work_start(lines: list[str]) -> int:
        try:
            return next(
                index + 1 for index, line in enumerate(lines) if line.startswith("Опыт работы")
            )
        except StopIteration as error:
            raise ValueError("В резюме не найден раздел опыта работы") from error

    def _find_work_end(self, lines: list[str], work_start: int) -> int:
        return next(
            (
                index
                for index in range(work_start, len(lines))
                if any(lines[index].startswith(section) for section in self._top_sections)
            ),
            len(lines),
        )

    def _find_narrative_start(self, lines: list[str], start: int, end: int) -> int:
        for index in range(start + 1, end):
            line = lines[index]
            if line.startswith("•"):
                continue
            if (
                line.startswith("-")
                or line in {"Обязанности", "Достижения", "Проекты:"}
                or self._project.search(line) is not None
                or len(line) >= 75
            ):
                return index
        raise ValueError(f"Не найдено описание места работы, начинающегося со строки {start + 1}")

    def _project_spans(
        self,
        lines: list[str],
        start: int,
        end: int,
    ) -> list[tuple[int, int, int, str]]:
        spans: list[tuple[int, int, int, str]] = []
        index = start
        while index < end:
            singular = self._project.search(lines[index])
            if singular is not None:
                next_marker = next(
                    (
                        following
                        for following in range(index + 1, end)
                        if self._project.search(lines[following]) is not None
                        or lines[following] == "Проекты:"
                    ),
                    end,
                )
                spans.append((index, index, next_marker, singular.group(1).strip()))
                index = next_marker
                continue
            if lines[index] == "Проекты:":
                collection_end = next(
                    (
                        following
                        for following in range(index + 1, end)
                        if re.match(r"^[СC]тек\b", lines[following], re.IGNORECASE) is not None
                    ),
                    end,
                )
                starts = [
                    following
                    for following in range(index + 1, collection_end)
                    if self._project_item.search(lines[following]) is not None
                ]
                if starts:
                    for position, project_start in enumerate(starts):
                        project_end = (
                            starts[position + 1] if position + 1 < len(starts) else collection_end
                        )
                        match = self._project_item.search(lines[project_start])
                        assert match is not None
                        spans.append((index, project_start, project_end, match.group(1).strip()))
                elif index + 1 < collection_end:
                    spans.append((index, index + 1, collection_end, "Проекты"))
                index = collection_end
                continue
            index += 1
        return spans

    def _work_label(
        self,
        lines: list[str],
        start: int,
        narrative_start: int,
        entry_number: int,
    ) -> str:
        candidates = []
        for line in lines[start:narrative_start]:
            if (
                self._date_line.search(line) is not None
                or self._duration.search(line) is not None
                or line.casefold() == "настоящее время"
                or line.startswith("•")
                or re.fullmatch(r"https?://\S+|[\w.-]+\.[a-zа-я]{2,}", line, re.IGNORECASE)
            ):
                continue
            candidates.append(line)
        if len(candidates) >= 2:
            return f"{candidates[0]} — {candidates[-1]}"
        if candidates:
            return candidates[-1]
        return f"Место работы {entry_number}"

    @staticmethod
    def _block(
        existing: list[ResumeNarrativeBlock],
        kind: ResumeBlockKind,
        label: str,
        start: int,
        end: int,
        lines: list[str],
    ) -> ResumeNarrativeBlock:
        source_text = "\n".join(lines[start:end]).strip()
        if not source_text:
            raise ValueError(f"Пустой блок резюме: {label}")
        return ResumeNarrativeBlock(
            index=len(existing) + 1,
            kind=kind,
            label=label,
            start_line=start,
            end_line=end,
            source_text=source_text,
        )


class ResumeImprovementService:
    def __init__(self, session: Session, data_dir: Path, model: ResumeTextModel) -> None:
        self._session = session
        self._data_dir = data_dir
        self._model = model
        self._extractor = ResumeBlockExtractor()

    def improve(
        self,
        account_id: int,
        answer_provider: AnswerProvider,
        *,
        target_role: str | None = None,
        vacancy_limit: int = 50,
    ) -> ResumeImprovementResult:
        if not 1 <= vacancy_limit <= 200:
            raise ValueError("Размер выборки вакансий должен быть от 1 до 200")
        profile, resume = self._active_resume(account_id)
        source_content = (resume.content_text or "").strip()
        if not source_content:
            raise ValueError("У активного резюме нет разобранного текста")
        role = (target_role or resume.title).strip()
        if not role:
            raise ValueError("Нужно указать направление поиска")

        structure = self._extractor.extract(source_content)
        vacancy_context = self._vacancy_context(account_id, vacancy_limit)
        improved: list[ImprovedResumeBlock] = []
        for block in structure.blocks:
            context = ResumePromptContext(
                kind=block.kind,
                source_block=block.source_text,
                target_role=role,
                vacancy_context=vacancy_context,
            )
            assessments = self._assess_questions(context)
            questions = select_missing_questions(assessments)
            answers = tuple(
                self._answer(block, question, answer_provider) for question in questions
            )
            rewrite_response = self._model.complete(
                SYSTEM_PROMPT,
                build_rewrite_prompt(
                    context,
                    tuple(
                        ResumeQuestionAnswer(question, answer)
                        for question, answer in zip(questions, answers, strict=True)
                    ),
                ),
            )
            improved_text = self._normalize_model_text(rewrite_response)
            improved.append(
                ImprovedResumeBlock(
                    index=block.index,
                    kind=block.kind,
                    label=block.label,
                    source_text=block.source_text,
                    improved_text=improved_text,
                    questions=questions,
                    answers=answers,
                )
            )

        improved_blocks = tuple(improved)
        improved_content = self._extractor.assemble(structure, improved_blocks)
        draft_path, report_path = self._write_result(
            account_id,
            resume,
            role,
            improved_content,
            improved_blocks,
        )
        source_unchanged = resume.content_text == source_content or (
            resume.content_text is not None and resume.content_text.strip() == source_content
        )
        if profile.active_resume_id != resume.id:
            source_unchanged = False
        return ResumeImprovementResult(
            resume_id=resume.id,
            target_role=role,
            model_name=self._model.model_name,
            draft_path=draft_path,
            report_path=report_path,
            blocks=improved_blocks,
            source_unchanged=source_unchanged,
        )

    def _assess_questions(
        self,
        context: ResumePromptContext,
    ) -> tuple[ResumeQuestionAssessment, ...]:
        prompt = build_questions_prompt(context)
        response = self._model.complete(SYSTEM_PROMPT, prompt)
        try:
            return parse_question_assessments(response)
        except ValueError:
            retry_prompt = (
                f"{prompt}\n\n"
                "Предыдущий ответ имел неверный формат. Верни только корректный JSON-массив "
                "из пяти объектов без разметки и пояснений."
            )
            return parse_question_assessments(self._model.complete(SYSTEM_PROMPT, retry_prompt))

    def _active_resume(self, account_id: int) -> tuple[CandidateProfileModel, ResumeModel]:
        profile = self._session.scalar(
            select(CandidateProfileModel).where(CandidateProfileModel.account_id == account_id)
        )
        if profile is None or profile.active_resume_id is None:
            raise LookupError("Активное резюме не выбрано; сначала импортируйте его")
        resume = self._session.get(ResumeModel, profile.active_resume_id)
        if resume is None or resume.account_id != account_id:
            raise LookupError("Активное резюме не найдено")
        return profile, resume

    def _vacancy_context(self, account_id: int, vacancy_limit: int) -> str:
        vacancies = self._session.scalars(
            select(VacancyModel)
            .join(DirectionVacancyModel)
            .join(CareerDirectionModel)
            .where(
                CareerDirectionModel.account_id == account_id,
                CareerDirectionModel.is_active.is_(True),
            )
            .order_by(VacancyModel.published_at.desc().nullslast(), VacancyModel.id.desc())
            .limit(vacancy_limit)
        ).unique()
        counts: Counter[str] = Counter()
        display: dict[str, str] = {}
        vacancy_count = 0
        for vacancy in vacancies:
            vacancy_count += 1
            for raw_skill in vacancy.key_skills:
                skill = " ".join(raw_skill.split()).strip(" ,.;")
                if not 2 <= len(skill) <= 80:
                    continue
                key = skill.casefold()
                counts[key] += 1
                display.setdefault(key, skill)
        if not counts:
            return "Дополнительных ориентиров по сохраненным вакансиям нет."
        top = sorted(counts, key=lambda key: (-counts[key], display[key].casefold()))[:15]
        return "\n".join(
            f"- {display[key]}: {counts[key]} из {vacancy_count} вакансий" for key in top
        )

    @staticmethod
    def _answer(
        block: ResumeNarrativeBlock,
        question: str,
        answer_provider: AnswerProvider,
    ) -> str:
        answer = answer_provider(block, question).strip()
        if not answer:
            raise ValueError(f"Не получен ответ для блока «{block.label}»")
        if len(answer) > 4000:
            raise ValueError("Ответ на уточняющий вопрос слишком длинный")
        return answer

    @staticmethod
    def _normalize_model_text(response: str) -> str:
        value = response.strip()
        fenced = re.fullmatch(
            r"```(?:markdown|text)?\s*(.*?)\s*```",
            value,
            flags=re.DOTALL | re.IGNORECASE,
        )
        if fenced is not None:
            value = fenced.group(1).strip()
        lines: list[str] = []
        for raw_line in value.splitlines():
            line = raw_line.strip()
            if line.startswith("* "):
                line = "- " + line[2:]
            if line.startswith("**") and line.endswith("**") and len(line) > 4:
                line = line[2:-2]
            lines.append(line)
        normalized = "\n".join(lines).strip()
        if not normalized:
            raise ValueError("Модель вернула пустой блок резюме")
        if len(normalized) > 20_000:
            raise ValueError("Модель вернула слишком длинный блок резюме")
        return normalized

    def _write_result(
        self,
        account_id: int,
        resume: ResumeModel,
        target_role: str,
        improved_content: str,
        blocks: tuple[ImprovedResumeBlock, ...],
    ) -> tuple[Path, Path]:
        output_dir = self._data_dir / "resume-improvements" / f"account-{account_id}"
        output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now(UTC).strftime("%Y%m%d-%H%M%S")
        digest = hashlib.sha256(improved_content.encode("utf-8")).hexdigest()[:10]
        base_name = f"resume-{resume.id}-improved-{timestamp}-{digest}"
        draft_path = output_dir / f"{base_name}.docx"
        report_path = output_dir / f"{base_name}.json"

        document = Document()
        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10)
        for raw_line in improved_content.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith(("Опыт работы", "Образование", "Навыки", "Дополнительная")):
                document.add_heading(line, level=1)
            elif line.rstrip(":") in {
                "Задачи и вклад",
                "Результаты",
                "Вклад",
                "Результат",
                "Технологии",
                "Назначение",
            }:
                document.add_heading(line, level=2)
            elif line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(line)

        temporary_docx: Path | None = None
        temporary_json: Path | None = None
        report = {
            "source_resume_id": resume.id,
            "source_sha256": resume.source_sha256,
            "target_role": target_role,
            "model_name": self._model.model_name,
            "question_prompt_version": QUESTION_PROMPT_VERSION,
            "rewrite_prompt_version": REWRITE_PROMPT_VERSION,
            "created_at": datetime.now(UTC).isoformat(),
            "blocks": [asdict(block) for block in blocks],
        }
        try:
            with NamedTemporaryFile(suffix=".docx", dir=output_dir, delete=False) as temporary:
                temporary_docx = Path(temporary.name)
            document.save(str(temporary_docx))
            os.replace(temporary_docx, draft_path)

            with NamedTemporaryFile(
                suffix=".json",
                dir=output_dir,
                delete=False,
                mode="w",
                encoding="utf-8",
            ) as temporary:
                temporary_json = Path(temporary.name)
                json.dump(report, temporary, ensure_ascii=False, indent=2, default=str)
            os.replace(temporary_json, report_path)
        finally:
            if temporary_docx is not None:
                temporary_docx.unlink(missing_ok=True)
            if temporary_json is not None:
                temporary_json.unlink(missing_ok=True)
        return draft_path, report_path
