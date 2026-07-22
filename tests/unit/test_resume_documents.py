from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as WordDocument

from hugin.adapters import resume_documents
from hugin.adapters.resume_documents import ResumeDocumentError, ResumeDocumentReader
from hugin.domain.resumes import ResumeSourceType
from hugin.services.resume_profile import ResumeProfileExtractor


def write_resume(path: Path) -> None:
    document = WordDocument()
    for line in (
        "Иван Иванов",
        "Проживает: Санкт-Петербург",
        "Гражданство: Россия, есть разрешение на работу: Россия",
        "Не готов к переезду, готов к редким командировкам",
        "Желаемая должность и зарплата",
        "Python backend разработчик",
        "Тип занятости: полная занятость",
        "Формат работы: удалённо, гибрид",
        "Опыт работы — 3 года",
        "Компания",
        "Backend-разработчик",
        "Разрабатывал API на Python и FastAPI, работал с PostgreSQL и Redis.",
        "Образование",
        "Высшее образование по прикладной информатике.",
        "Повышение квалификации, курсы",
        "Курс Python-разработчика.",
        "Знание языков Русский — Родной",
        "Английский — B1 — Средний",
        "Навыки",
        "Python FastAPI SQL Docker Git PostgreSQL Redis",
        "Опыт вождения",
        "Права категории B",
        "Дополнительная информация",
        "Пишу проверяемый код, использую тесты и контейнеры.",
        "GitHub: github.com/example",
        "user@example.com",
    ):
        document.add_paragraph(line)
    document.save(str(path))


def test_docx_is_read_and_profile_is_extracted(tmp_path: Path) -> None:
    path = tmp_path / "Резюме ИТ.docx"
    write_resume(path)

    resume = ResumeDocumentReader().read(path)
    profile = ResumeProfileExtractor().extract(resume)

    assert resume.source_type is ResumeSourceType.DOCX
    assert resume.original_name == "Резюме ИТ.docx"
    assert len(resume.sha256) == 64
    assert profile.display_name == "Иван Иванов"
    assert profile.title == "Python backend разработчик"
    assert {fact.category for fact in profile.facts} >= {
        "desired_position",
        "education",
        "skills",
        "work_experience",
    }
    missing = {question.key for question in profile.missing_questions}
    assert "salary_expectation" in missing
    assert "available_from" in missing
    assert "work_format" not in missing
    assert "english_level" not in missing
    assert "portfolio" not in missing


def test_repeated_inline_skills_heading_is_extracted() -> None:
    extractor = ResumeProfileExtractor()
    lines = [
        "Навыки",
        "Знание языков Русский — Родной",
        "Английский — B1 — Средний",
        "Навыки Python FastAPI SQL Docker Git PostgreSQL Redis",
        "Опыт вождения",
        "Права категории B",
    ]

    assert extractor._section(lines, "skills") == ("Python FastAPI SQL Docker Git PostgreSQL Redis")
    assert extractor._languages(lines) == "Русский — Родной\nАнглийский — B1 — Средний"


def test_pdf_with_cyrillic_name_is_read(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    path = tmp_path / "Резюме программиста.pdf"
    path.write_bytes(b"test-pdf")

    class FakePage:
        def extract_text(self) -> str:
            return "Желаемая должность и зарплата\nPython разработчик\n" + "Опыт Python. " * 30

    class FakeReader:
        def __init__(self) -> None:
            self.is_encrypted = False
            self.pages = [FakePage(), FakePage()]

    monkeypatch.setattr(resume_documents, "PdfReader", lambda _: FakeReader())

    resume = ResumeDocumentReader().read(path)

    assert resume.source_type is ResumeSourceType.PDF
    assert resume.page_count == 2
    assert "Python разработчик" in resume.text


def test_empty_and_unsupported_documents_are_rejected(tmp_path: Path) -> None:
    empty = tmp_path / "empty.docx"
    WordDocument().save(str(empty))
    unsupported = tmp_path / "resume.txt"
    unsupported.write_text("resume", encoding="utf-8")

    with pytest.raises(ResumeDocumentError, match="текстового слоя"):
        ResumeDocumentReader().read(empty)
    with pytest.raises(ResumeDocumentError, match="PDF и DOCX"):
        ResumeDocumentReader().read(unsupported)
